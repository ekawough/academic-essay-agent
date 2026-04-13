import os
import re
import anthropic
from typing import Optional

client = anthropic.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))

PAPER_CONFIGS = {
    "research_paper": {
        "word_count": "2000-3000",
        "sections": ["Abstract", "Introduction", "Literature Review", "Methodology", "Results & Discussion", "Conclusion", "References"],
        "description": "A focused academic research paper with clear thesis, evidence, and argumentation.",
        "tone": "formal academic, objective, evidence-based"
    },
    "bachelor": {
        "word_count": "3000-5000",
        "sections": ["Abstract", "Introduction", "Literature Review", "Theoretical Framework", "Methodology", "Analysis & Discussion", "Conclusion", "References"],
        "description": "A bachelor's level academic essay demonstrating understanding of the field, critical analysis, and academic writing conventions.",
        "tone": "formal academic, analytical, well-structured, demonstrates critical thinking"
    },
    "master": {
        "word_count": "5000-8000",
        "sections": ["Abstract", "Introduction", "Literature Review", "Theoretical Framework", "Methodology", "Findings & Analysis", "Discussion", "Conclusion", "References"],
        "description": "A master's level paper with sophisticated analysis, original argumentation, and deep engagement with scholarly literature.",
        "tone": "highly academic, nuanced, demonstrates advanced critical analysis and original thinking"
    }
}

async def write_essay(
    topic: str,
    paper_type: str = "bachelor",
    language: str = "en",
    context: str = None,
    additional_instructions: str = None
) -> dict:
    """
    Writes a full academic essay using Claude.
    Injects verified research context to prevent hallucination.
    """
    config = PAPER_CONFIGS.get(paper_type, PAPER_CONFIGS["bachelor"])

    system_prompt = f"""You are an expert academic ghostwriter specializing in {paper_type}-level essays.

Your writing standards:
- Tone: {config['tone']}
- Target length: {config['word_count']} words
- Structure: {', '.join(config['sections'])}
- Use proper in-text citations (Author, Year) format for ALL claims
- Every factual claim must be supported by a citation from the provided research context
- Do NOT invent citations, DOIs, or authors — only use sources from the research context provided
- Use formal academic language appropriate for university submission
- Include smooth transitions between sections
- Avoid passive voice overuse; balance active and passive
- Vary sentence structure and length for readability
- Demonstrate critical thinking, not just description
- Write a genuine thesis statement in the introduction

CRITICAL RULE: If you reference a source, it MUST exist in the research context provided. 
If no context is provided for a claim, write "further research is needed" or omit the claim.

Output format:
Return the full essay with clear section headers using ## for each section.
At the very end, include a "## References" section listing all cited sources in APA 7th edition format.
"""

    context_section = ""
    if context:
        context_section = f"""
## Verified Research Context (use ONLY these sources for citations):

{context[:8000]}

---
"""

    instructions_section = ""
    if additional_instructions:
        instructions_section = f"\nAdditional requirements: {additional_instructions}\n"

    user_prompt = f"""Write a {config['description']}

Topic: {topic}
Paper type: {paper_type}
Language: {language}
{instructions_section}
{context_section}

Write the complete essay now. Begin with the Abstract section.
"""

    response = client.messages.create(
        model="claude-opus-4-5",
        max_tokens=8000,
        messages=[
            {"role": "user", "content": user_prompt}
        ],
        system=system_prompt
    )

    full_text = response.content[0].text

    # Parse sections
    sections = _parse_sections(full_text)
    citations = _extract_citations(full_text)
    word_count = len(full_text.split())

    # Extract title from intro if present
    title = _extract_title(full_text, topic)

    return {
        "title": title,
        "content": full_text,
        "sections": sections,
        "citations": citations,
        "word_count": word_count,
        "paper_type": paper_type,
        "model_used": "claude-opus-4-5"
    }

def _parse_sections(text: str) -> list:
    """Extract section names from the essay."""
    sections = re.findall(r'^##\s+(.+)$', text, re.MULTILINE)
    return sections

def _extract_citations(text: str) -> list:
    """Extract in-text citations and reference list."""
    # Find references section
    ref_match = re.search(r'##\s+References?\s*\n([\s\S]+?)(?:##|$)', text, re.IGNORECASE)
    if ref_match:
        ref_block = ref_match.group(1).strip()
        references = [line.strip() for line in ref_block.split('\n') if line.strip() and not line.strip().startswith('#')]
        return references
    return []

def _extract_title(text: str, fallback_topic: str) -> str:
    """Try to extract a proper title from the essay."""
    # Look for a title at the top before Abstract
    lines = text.split('\n')
    for line in lines[:10]:
        line = line.strip()
        if line and not line.startswith('#') and len(line) > 20 and len(line) < 150:
            if not any(kw in line.lower() for kw in ['abstract', 'introduction', 'write', 'here']):
                return line
    # Check for # Title at top
    title_match = re.search(r'^#\s+(.+)$', text, re.MULTILINE)
    if title_match:
        return title_match.group(1)
    return f"Academic Essay: {fallback_topic.title()}"

def export_essay_docx(title: str, content: str, citations: list) -> bytes:
    """Export essay to DOCX format."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import io

        doc = Document()

        # Set margins
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1.25)

        # Title
        title_para = doc.add_heading(title, 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        # Process content
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
            if line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('**') and line.endswith('**'):
                p = doc.add_paragraph()
                run = p.add_run(line[2:-2])
                run.bold = True
            else:
                doc.add_paragraph(line)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()

    except ImportError:
        return content.encode('utf-8')
